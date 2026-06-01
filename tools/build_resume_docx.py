from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path("output/resume")
OUT_DOCX = OUT_DIR / "王玉隆-Java后端开发实习生-简历.docx"

NAVY = "2F3B4A"
LIGHT = "F1F3F5"
TEXT = "4A4A4A"
MUTED = "666666"
BORDER = "2F3B4A"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_table_borders(table, color="FFFFFF", size="0"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil" if size == "0" else "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=50, start=90, bottom=50, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = pg_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            pg_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "12")
        element.set(qn("w:space"), "18")
        element.set(qn("w:color"), BORDER)


def set_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=9, bold=False, color=TEXT, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para(cell_or_doc, text="", size=9, bold=False, color=TEXT, align=None, before=0, after=2, font="Microsoft YaHei"):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, font=font)
    return p


def clear_cell(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def section_label(doc, title):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    set_width(table.cell(0, 0), 3.7)
    set_width(table.cell(0, 1), 14.3)
    left, right = table.cell(0, 0), table.cell(0, 1)
    for cell in (left, right):
        set_cell_margins(cell, top=55, bottom=55, start=110, end=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(left, NAVY)
    set_cell_shading(right, LIGHT)
    clear_cell(left)
    clear_cell(right)
    p = para(left, title, size=12, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    p.paragraph_format.line_spacing = 1.0
    return table


def add_numbered_lines(doc, items, left_indent_cm=0.45):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(left_indent_cm)
        p.paragraph_format.first_line_indent = Cm(-0.05)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r1 = p.add_run(f"{idx}. ")
        set_run_font(r1, 9, color=MUTED)
        r2 = p.add_run(item)
        set_run_font(r2, 9, color=TEXT)


def add_bullets(doc, items, indent_cm=1.0):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.06
        r1 = p.add_run("o ")
        set_run_font(r1, 8.6, color=MUTED)
        r2 = p.add_run(item)
        set_run_font(r2, 8.6, color=TEXT)


def add_kv_line(doc, key, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r1 = p.add_run(f"{key}：")
    set_run_font(r1, 8.8, bold=True, color=TEXT)
    r2 = p.add_run(value)
    set_run_font(r2, 8.8, color=TEXT)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.72)
    section.bottom_margin = Cm(0.72)
    section.left_margin = Cm(0.82)
    section.right_margin = Cm(0.82)
    set_page_border(section)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9)

    # Decorative top tab, matching the supplied reference style.
    tab = doc.add_table(rows=1, cols=3)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.autofit = False
    set_table_borders(tab)
    widths = [6.1, 5.0, 6.1]
    for i, w in enumerate(widths):
        set_width(tab.cell(0, i), w)
        set_cell_margins(tab.cell(0, i), top=15, bottom=15, start=0, end=0)
        clear_cell(tab.cell(0, i))
    set_cell_shading(tab.cell(0, 1), NAVY)
    p = para(tab.cell(0, 1), "  ", size=7, color="FFFFFF", after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    set_table_borders(header)
    left, right = header.cell(0, 0), header.cell(0, 1)
    set_width(left, 4.3)
    set_width(right, 13.5)
    for cell in (left, right):
        clear_cell(cell)
        set_cell_margins(cell, top=70, bottom=80, start=130, end=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    photo_table = left.add_table(rows=1, cols=1)
    photo_table.autofit = False
    set_width(photo_table.cell(0, 0), 3.1)
    set_table_borders(photo_table, color=NAVY, size="10")
    photo = photo_table.cell(0, 0)
    set_cell_margins(photo, top=620, bottom=620, start=80, end=80)
    clear_cell(photo)
    p = para(photo, "证件照", size=9, bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    p.paragraph_format.line_spacing = 1.0
    # Word requires a table cell that contains a nested table to still end with a paragraph.
    tail = left.add_paragraph()
    tail.paragraph_format.space_after = Pt(0)

    para(right, "王玉隆 | Java 后端开发实习生", size=15, bold=True, color=TEXT, after=8)
    para(right, "电话：182-3862-1630", size=11, color=MUTED, after=5)
    para(right, "邮箱：3295791762@qq.com", size=11, color=MUTED, after=5)
    para(right, "求职意向：Java 后端开发 / 后端开发实习", size=10, color=MUTED, after=0)

    section_label(doc, "专业技能")
    add_numbered_lines(doc, [
        "Java：熟悉 Java 基础、面向对象、集合框架、异常处理，了解多线程基础使用场景。",
        "Spring Boot / Spring MVC：熟悉分层开发、RESTful 接口、参数校验、全局异常处理和统一响应封装。",
        "MyBatis / MyBatis-Plus / MySQL：了解 ORM 原理，能完成常见 CRUD、Mapper 设计、索引与事务相关开发。",
        "Redis / SSE：了解 Redis 热点数据缓存场景，掌握 Server-Sent Events 流式推送在前后端交互中的使用。",
        "AI 工程实践：了解 Agent Workflow、Tool Calling、RAG 检索、结构化 JSON 输出和学习记忆设计。",
        "其他：熟悉 Git、Maven、Linux 常用命令，了解 Next.js、Tailwind CSS、Docker Compose 基础使用。",
    ])

    section_label(doc, "项目经历")
    add_kv_line(doc, "项目名称", "AI Interview Coach Agent：基于 Agent 工作流的 Java 面试训练与代码诊断系统")
    add_kv_line(doc, "项目时间", "2026.05 - 至今")
    add_kv_line(doc, "项目技术", "Spring Boot 3、Java 17、MyBatis-Plus、MySQL、Redis、SSE、Next.js 14、Piston API、大模型 API")
    add_kv_line(
        doc,
        "项目描述",
        "面向 Java 后端面试训练的代码诊断系统，围绕“代码提交 -> 执行观察 -> RAG 检索 -> AI 诊断/代码评审 -> 弱点记忆 -> 训练计划”构建可解释 Agent 闭环。",
    )
    add_kv_line(doc, "个人职责", "")
    add_bullets(doc, [
        "参与后端核心链路开发，按 Controller、Service、Agent Tool、Mapper 分层组织提交、判题、诊断和训练计划逻辑。",
        "封装 Java Solution 模式代码执行流程，保留用户原始代码，并通过 CodeWrapper 适配 Two Sum、反转链表、买卖股票等 Hot100 题型。",
        "设计 Agent 工作流：Planner -> CodeExecutionTool -> Observation -> RagRetrieveTool -> ErrorClassifierTool / CodeReviewTool -> WeaknessTrackerTool -> TrainingPlannerTool。",
        "实现 MySQL 结构化 RAG V1，索引题目、知识卡片、历史诊断和错题卡，并在检索条件中隔离 user_id，避免用户学习记忆串读。",
        "实现 SSE 诊断流，前端通过 ReadableStream 展示 AgentStep，支持代码执行、RAG 检索、错误分类、弱点更新、训练计划等步骤实时展示。",
        "建设学习记忆闭环：失败提交生成 AI 诊断、弱点事件、错题卡和训练计划；Accepted 提交进入轻量代码评审分支，输出复杂度和面试表达建议。",
        "使用 Redis 缓存题目列表、题目详情和代码模板，MySQL 负责提交、诊断、弱点、错题卡、训练计划等持久化学习数据。",
    ], indent_cm=0.95)

    section_label(doc, "工作经历")
    work = doc.add_table(rows=1, cols=3)
    work.autofit = False
    set_table_borders(work)
    for i, w in enumerate([5.0, 7.0, 5.6]):
        set_width(work.cell(0, i), w)
        clear_cell(work.cell(0, i))
        set_cell_margins(work.cell(0, i), top=30, bottom=25, start=100, end=100)
    para(work.cell(0, 0), "2025.06 - 2025.08", size=9, bold=True, color=TEXT, after=0)
    para(work.cell(0, 1), "科大讯飞", size=9, bold=True, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(work.cell(0, 2), "产品运营实习生", size=9, bold=True, color=TEXT, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)
    add_bullets(doc, [
        "使用公司内部后台管理系统，负责数据审核、修改与录入，保证业务数据准确性。",
        "收集教师与学生在系统使用过程中的问题，整理登录、权限、课程配置等反馈并同步技术团队。",
        "在实际业务场景中理解后台系统功能设计、业务流程和跨角色协作方式。",
    ], indent_cm=0.75)

    section_label(doc, "教育经历")
    edu = doc.add_table(rows=1, cols=4)
    edu.autofit = False
    set_table_borders(edu)
    for i, w in enumerate([4.8, 4.5, 5.6, 2.8]):
        set_width(edu.cell(0, i), w)
        clear_cell(edu.cell(0, i))
        set_cell_margins(edu.cell(0, i), top=45, bottom=40, start=90, end=90)
    para(edu.cell(0, 0), "2023.09 - 2027.06", size=9, bold=True, color=TEXT, after=0)
    para(edu.cell(0, 1), "信阳学院", size=9, bold=True, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(edu.cell(0, 2), "计算机科学与技术", size=9, bold=True, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(edu.cell(0, 3), "本科", size=9, bold=True, color=TEXT, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)

    doc.save(OUT_DOCX)
    print(OUT_DOCX.resolve())


if __name__ == "__main__":
    build()
